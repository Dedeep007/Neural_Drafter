from docx import Document

def extract_first_table_as_dict(doc_path):
    # Load the Word document
    doc = Document(doc_path)
    
    # Ensure there's at least one table
    if not doc.tables:
        return []

    # Get the first table
    table = doc.tables[0]
    
    # Assume first row is the header
    keys = [cell.text.strip() for cell in table.rows[0].cells]
    
    # Extract remaining rows as dictionary entries
    data = []
    for row in table.rows[1:]:
        values = [cell.text.strip() for cell in row.cells]
        # Make dictionary from keys and values
        row_dict = dict(zip(keys, values))
        data.append(row_dict)
    
    return data

if __name__ == "__main__":
    doc_path = "output.docx"
    table_data = extract_first_table_as_dict(doc_path)
    print(table_data)