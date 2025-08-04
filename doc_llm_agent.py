from gradio_client import Client, handle_file
from docx import Document
from docx.shared import Inches
import ast
import re
from dotenv import load_dotenv
from langchain.tools import Tool
from langchain_experimental.tools import PythonREPLTool
from langchain.agents import initialize_agent
from langchain_groq import ChatGroq
import tempfile
from PIL import Image
import os

load_dotenv()

 # Gradio OCR Tool
def ocr_tool(image_path: str) -> str:
    try:
        client = Client("IotaCluster/OCR")
        # Accept both local file paths and URLs
        image_input = image_path
        if not (image_path.startswith('http://') or image_path.startswith('https://')):
            image_input = handle_file(image_path)
        else:
            image_input = handle_file(image_path)  # handle_file also works with URLs
        result = client.predict(
            image=image_input,
            language=["eng"],
            api_name="/predict"
        )
        return str(result)
    except Exception as e:
        return f"OCR failed: {str(e)}"

# Word Document Builder
class DocumentBuilder:
    def __init__(self):
        self.doc = Document()

    def add_image_with_text(self, image_path, matched_dict):
        self.doc.add_picture(image_path, width=Inches(5))
        if matched_dict:
            # Add a table with keys as headers and values as a row, and set borders
            table = self.doc.add_table(rows=2, cols=len(matched_dict))
            table.style = 'Table Grid'  # Adds borders to the table
            hdr_cells = table.rows[0].cells
            val_cells = table.rows[1].cells
            for i, (k, v) in enumerate(matched_dict.items()):
                hdr_cells[i].text = str(k)
                val_cells[i].text = str(v)
        self.doc.add_paragraph("\n")

    def save(self, path="output.docx"):
        self.doc.save(path)
        return path

# Core Function with LLM Matching

def compile_images_and_text_to_doc(image_objs, text_objects):
    """
    image_objs: list of PIL.Image.Image objects
    text_objects: list of dicts
    """
    ocr_tool_instance = Tool(
        name="ImageOCR",
        func=ocr_tool,
        description="Extracts text content from an image using OCR via Gradio"
    )

    doc_builder = DocumentBuilder()

    llm = ChatGroq(temperature=0, model_name="qwen/qwen3-32b")
    agent = initialize_agent(
        tools=[ocr_tool_instance, PythonREPLTool()],
        llm=llm,
        agent_type="chat-zero-shot-react-description",
        verbose=True,
        handle_parsing_errors=True
    )

    temp_files = []
    try:
        for img in image_objs:
            # Save PIL image to a temporary file
            temp_file = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            img.save(temp_file, format="PNG")
            temp_file.close()
            temp_files.append(temp_file.name)
            image_path = temp_file.name

            print(f"🔍 Processing: {image_path}")
            ocr_text = agent.run(f"Use ImageOCR to extract text from image: {image_path}")

            prompt = (
                f"You are a helpful AI. Here is some OCR text extracted from an image:\n{ocr_text}\n\n"
                f"Below is a list of product entries (each as a dictionary).\n{text_objects}\n"
                f"Your task is to match the OCR text with the most relevant dictionary based on the fields."
                f" Return only the best matching dictionary in raw Python dictionary format."
            )

            llm_match = agent.run(prompt)

            # Extract the first dictionary from the LLM output using regex
            dict_match = re.search(r'\{[^\{\}]*\}', llm_match, re.DOTALL)
            if dict_match:
                dict_str = dict_match.group(0)
                try:
                    matched_dict = ast.literal_eval(dict_str)
                except Exception:
                    matched_dict = {"Match": dict_str}
            else:
                matched_dict = {"Match": llm_match.strip()[:200]}

            doc_builder.add_image_with_text(image_path, matched_dict)

        path = doc_builder.save()
        print(f"✅ Final document saved at: {path}")
    finally:
        # Clean up temp files
        for f in temp_files:
            try:
                os.remove(f)
            except Exception:
                pass

# Example usage:
if __name__ == "__main__":
    from PIL import Image
    image_objs = [Image.open("4090.png"), Image.open("eg_img.png"), Image.open("5080.png")]
    text_objects = [
        {"Name": "RTX 5080 GPU", "Cost": "126999", "Quantity": 3},
        {"Name": "RTX 4090", "Cost": "292999", "Quantity": 5, "Description": "High-end GPU for gaming and AI"},
        {"Name": "groq", "Cost": "N/A", "Quantity": "N/A", "Description": "Service provider for AI models"},
        {"Name": "T 1000", "Cost": "N/A", "Quantity": "N/A", "Description": "High-end GPU for gaming and AI"}
    ]
    compile_images_and_text_to_doc(image_objs, text_objects)
